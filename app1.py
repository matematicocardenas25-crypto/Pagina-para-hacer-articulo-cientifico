import streamlit as st
import numpy as np
from docx import Document
from docx.shared import Pt
from PIL import Image
import io
import easyocr
import datetime

# --- CONFIGURACIÓN ---
st.set_page_config(page_title="Redactor Científico Prof. Cárdenas", layout="wide")

# --- FUNCIONES DE GENERACIÓN ---
def generar_word_articulo(datos, bibliografia):
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Arial'
    style.font.size = Pt(11)
    
    # Título
    titulo = doc.add_heading(datos['titulo'], 0)
    titulo.alignment = 1 # Centrado
    
    # Resúmenes
    doc.add_heading('RESUMEN', level=1)
    doc.add_paragraph(datos['resumen_es'])
    
    doc.add_heading('ABSTRACT', level=1)
    doc.add_paragraph(datos['resumen_en'])
    
    secciones = [
        ("INTRODUCCIÓN", "La presente investigación aborda el fenómeno desde una perspectiva analítica..."),
        ("METODOLOGÍA", f"Se aplicó un enfoque {datos['metodologia']}."),
        ("RESULTADOS", datos['cuerpo']),
        ("CONCLUSIONES", "Los resultados sugieren una correlación significativa entre las variables estudiadas.")
    ]
    
    for tit, cont in secciones:
        doc.add_heading(tit, level=1)
        doc.add_paragraph(cont)
    
    if bibliografia:
        doc.add_heading("BIBLIOGRAFÍA (Normas APA)", level=1)
        for cita in sorted(bibliografia):
            doc.add_paragraph(cita, style='List Bullet')
            
    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf

def generar_latex_articulo(datos, bibliografia):
    bib_items = "\n".join([f"\\item {c}" for c in sorted(bibliografia)])
    latex = f"""\\documentclass[12pt,a4paper]{{article}}
\\usepackage[utf8]{{inputenc}}
\\usepackage[spanish,english]{{babel}}
\\title{{{datos['titulo']}}}
\\author{{Prof. Ismael Antonio Cárdenas López}}

\\begin{{document}}
\\maketitle

\\selectlanguage{{spanish}}
\\begin{{abstract}}
{datos['resumen_es']}
\\end{{abstract}}

\\selectlanguage{{english}}
\\renewcommand{{\\abstractname}}{{Abstract}}
\\begin{{abstract}}
{datos['resumen_en']}
\\end{{abstract}}

\\selectlanguage{{spanish}}
\\section{{Introducción}}
Texto introductorio...
\\section{{Metodología}}
{datos['metodologia']}
\\section{{Resultados}}
{datos['cuerpo']}
\\section{{Bibliografía}}
\\begin{{itemize}}
{bib_items}
\\end{{itemize}}
\\end{{document}}"""
    return latex

# --- INTERFAZ DE USUARIO ---
st.title("📝 Redactor Académico Multilingüe")
st.info("Este módulo genera artículos científicos con Resumen (ES) y Abstract (EN) automáticos.")

# 1. OCR PARA IMÁGENES
with st.expander("📷 Extraer texto de imagen (OCR)"):
    archivo_img = st.file_uploader("Sube una foto del artículo", type=['jpg','png','jpeg'])
    texto_extraido = ""
    if archivo_img:
        reader = easyocr.Reader(['es', 'en'])
        img = Image.open(archivo_img)
        texto_extraido = "\n".join(reader.readtext(np.array(img), detail=0))
        st.text_area("Texto detectado:", value=texto_extraido, height=100)

# 2. FORMULARIO PRINCIPAL
with st.form("main_form"):
    titulo = st.text_input("Título del Artículo", "Análisis Estadístico sobre...")
    metodo = st.selectbox("Metodología", ["Cuantitativa", "Cualitativa", "Mixta"])
    
    col_es, col_en = st.columns(2)
    res_es = col_es.text_area("Resumen (Español)", height=150, placeholder="Escribe el resumen aquí...")
    res_en = col_en.text_area("Abstract (Inglés)", height=150, placeholder="Write the abstract here...")
    
    cuerpo = st.text_area("Contenido / Resultados", value=texto_extraido, height=200)
    
    btn_guardar = st.form_submit_button("✅ Procesar Artículo")

# 3. GESTOR DE BIBLIOGRAFÍA
st.subheader("📚 Referencias APA")
if 'bibliografia' not in st.session_state: st.session_state.bibliografia = []

with st.expander("Agregar Referencia"):
    c1, c2, c3 = st.columns([3,1,2])
    aut = c1.text_input("Autor")
    yr = c2.text_input("Año")
    ti = c3.text_input("Título")
    if st.button("Añadir"):
        st.session_state.bibliografia.append(f"{aut} ({yr}). {ti}.")
        st.rerun()

# 4. BOTONES DE DESCARGA
if btn_guardar:
    datos = {
        "titulo": titulo, "metodologia": metodo, 
        "resumen_es": res_es, "resumen_en": res_en, "cuerpo": cuerpo
    }
    
    st.success("¡Documento estructurado con éxito!")
    col_d1, col_d2 = st.columns(2)
    
    # Word
    buf_word = generar_word_articulo(datos, st.session_state.bibliografia)
    col_d1.download_button("📥 Descargar Word (.docx)", buf_word, f"{titulo}.docx")
    
    # LaTeX
    code_latex = generar_latex_articulo(datos, st.session_state.bibliografia)
    col_d2.download_button("📥 Descargar LaTeX (.tex)", code_latex.encode('utf-8'), f"{titulo}.tex")
